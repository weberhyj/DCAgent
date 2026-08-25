from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from fastapi.testclient import TestClient

from app.ingestion import KnowledgeIndexUnavailableError, KnowledgeIngestionQueue
from app.main import create_app
from app.repository import InMemoryChatRepository
from app.seed import build_seed_state
from app.text_parser import parse_knowledge_file, parse_knowledge_file_result
from app.word_facts import WordFactualIntent, normalize_fact_key


class KnowledgeIngestionPipelineTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = tempfile.TemporaryDirectory()
        self.repository = InMemoryChatRepository(build_seed_state())
        self.queue = KnowledgeIngestionQueue(self.repository)
        self.client = TestClient(
            create_app(
                repository=self.repository,
                upload_dir=Path(self.temp_dir.name),
                ingestion_queue=self.queue,
            )
        )

    def tearDown(self) -> None:
        self.temp_dir.cleanup()

    def test_upload_schedules_indexing_after_returning(self) -> None:
        body = ("第一段现金流分析。" * 80).encode("utf-8")

        response = self.client.post(
            "/api/knowledge/uploads",
            data={"classification": "内部·机密"},
            files={"file": ("cashflow-note.txt", body, "text/plain")},
        )

        self.assertEqual(response.status_code, 202)
        uploaded = response.json()[0]
        self.assertEqual(uploaded["name"], "cashflow-note.txt")
        self.assertEqual(uploaded["status"], "解析中")
        self.assertEqual(uploaded["records"], 0)

        sources = self.client.get("/api/knowledge/sources").json()
        indexed = sources[0]
        self.assertEqual(indexed["id"], uploaded["id"])
        self.assertEqual(indexed["status"], "已索引")
        self.assertGreater(indexed["records"], 1)

        chunks_response = self.client.get(f"/api/knowledge/sources/{uploaded['id']}/chunks")
        self.assertEqual(chunks_response.status_code, 200)
        chunks = chunks_response.json()
        self.assertEqual(len(chunks), indexed["records"])
        self.assertEqual(chunks[0]["sourceId"], uploaded["id"])
        self.assertEqual(chunks[0]["chunkIndex"], 0)
        self.assertIn("现金流", chunks[0]["text"])

    def test_parser_splits_long_text_file_into_ordered_chunks(self) -> None:
        path = Path(self.temp_dir.name) / "risk-note.md"
        path.write_text(
            "\n".join([f"风险提示 {index}: 回款周期变化" for index in range(120)]), encoding="utf-8"
        )

        chunks = parse_knowledge_file(path, source_id="kb-parser", source_type="文档")

        self.assertGreater(len(chunks), 1)
        self.assertEqual([chunk.chunk_index for chunk in chunks], list(range(len(chunks))))
        self.assertTrue(all(chunk.source_id == "kb-parser" for chunk in chunks))
        self.assertIn("风险提示", chunks[0].text)

    def test_parser_removes_nul_bytes_before_chunks_are_indexed(self) -> None:
        path = Path(self.temp_dir.name) / "nul-note.txt"
        path.write_bytes("差旅制度".encode() + b"\x00" + "审批流程".encode())

        chunks = parse_knowledge_file(path, source_id="kb-nul", source_type="文档")

        self.assertGreater(len(chunks), 0)
        self.assertTrue(all("\x00" not in chunk.text for chunk in chunks))
        self.assertIn("差旅制度", chunks[0].text)
        self.assertIn("审批流程", chunks[0].text)

    def test_non_docx_parse_result_keeps_chunk_compatibility_without_facts(self) -> None:
        path = Path(self.temp_dir.name) / "plain-note.txt"
        path.write_text("plain searchable passage", encoding="utf-8")

        result = parse_knowledge_file_result(path, source_id="kb-plain", source_type="TXT")

        self.assertIsInstance(result.chunks, tuple)
        self.assertEqual(result.facts, ())
        self.assertEqual([chunk.text for chunk in result.chunks], ["plain searchable passage"])

    def test_csv_chunks_keep_header_with_each_data_section(self) -> None:
        path = Path(self.temp_dir.name) / "sales.csv"
        path.write_text(
            "地区,日期,销售额\n华东,2025-01-01,100\n华东,2025-01-02,200\n",
            encoding="utf-8",
        )

        chunks = parse_knowledge_file(path, source_id="kb-csv", source_type="CSV")

        self.assertEqual(len(chunks), 1)
        self.assertIn("[CSV]", chunks[0].text)
        self.assertIn("地区 | 日期 | 销售额", chunks[0].text)
        self.assertIn("华东 | 2025-01-02 | 200", chunks[0].text)

    def test_headerless_csv_keeps_first_data_row(self) -> None:
        path = Path(self.temp_dir.name) / "headerless.csv"
        path.write_text(
            "华东,2025-01-01,100\n华东,2025-01-02,200\n",
            encoding="utf-8",
        )

        chunks = parse_knowledge_file(path, source_id="kb-headerless", source_type="CSV")

        self.assertEqual(len(chunks), 1)
        self.assertIn("数据行（无表头）", chunks[0].text)
        self.assertIn("华东 | 2025-01-01 | 100", chunks[0].text)
        self.assertIn("华东 | 2025-01-02 | 200", chunks[0].text)

    def test_xlsx_chunks_repeat_sheet_and_header_without_splitting_rows(self) -> None:
        path = Path(self.temp_dir.name) / "regional-sales.xlsx"
        from openpyxl import Workbook

        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "销售明细"
        worksheet.append(["地区", "日期", "销售额"])
        rows = []
        for index in range(80):
            row = ["华东", f"2025-01-{(index % 28) + 1:02d}", index]
            rows.append(" | ".join(str(value) for value in row))
            worksheet.append(row)
        workbook.save(path)

        chunks = parse_knowledge_file(path, source_id="kb-xlsx", source_type="XLSX")

        self.assertGreater(len(chunks), 1)
        for chunk in chunks:
            self.assertTrue(chunk.text.startswith("[销售明细]\n地区 | 日期 | 销售额\n"))
            self.assertLessEqual(len(chunk.text), 600)
            data_lines = chunk.text.splitlines()[2:]
            self.assertTrue(all(line in rows for line in data_lines))

    def test_tabular_chunks_cap_oversized_header_and_row(self) -> None:
        from openpyxl import Workbook

        path = Path(self.temp_dir.name) / "wide.xlsx"
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "宽表"
        worksheet.append(["列头" * 200])
        worksheet.append(["数据" * 500])
        workbook.save(path)

        chunks = parse_knowledge_file(path, source_id="kb-wide", source_type="XLSX")

        self.assertGreater(len(chunks), 1)
        self.assertTrue(all(len(chunk.text) <= 600 for chunk in chunks))

    def test_new_document_updates_postgres_then_active_qdrant_collection(self) -> None:
        events = []

        class Lifecycle:
            def upsert_source(inner_self, source_id, *, finalize=None, on_failure=None):
                events.append(("qdrant", len(self.repository.list_knowledge_chunks(source_id))))
                if finalize is None:
                    events.append("missing-finalize")
                    return "publication-1", 1
                indexed = finalize(
                    SimpleNamespace(
                        publication_id="publication-1",
                        indexed_point_count=1,
                    )
                )
                events.append("fence-release")
                return indexed

        source_id = "source-lifecycle"
        path = Path(self.temp_dir.name) / "lifecycle.txt"
        path.write_text("one searchable passage", encoding="utf-8")
        self.repository.add_uploaded_knowledge_source(
            source_id, "lifecycle.txt", "TXT", "internal", 0, str(path), 22, "text/plain"
        )
        original_complete = self.repository.complete_retrieval_source_indexing

        def complete_retrieval_index(*args, **kwargs):
            events.append("postgres-indexed")
            return original_complete(*args, **kwargs)

        self.repository.complete_retrieval_source_indexing = complete_retrieval_index
        queue = KnowledgeIngestionQueue(self.repository, index_lifecycle=Lifecycle())

        queue.enqueue(source_id, path, "TXT")
        queue.drain()

        self.assertEqual(events, [("qdrant", 1), "postgres-indexed", "fence-release"])
        self.assertEqual(self.repository.get_source_index_status(source_id), "indexed")

    def test_first_document_ensures_publication_after_empty_startup(self) -> None:
        events: list[str] = []

        class FreshLifecycle:
            def ensure_active_publication(inner_self):
                events.append("bootstrap-publication")

            def upsert_source(inner_self, source_id, *, finalize=None, on_failure=None):
                events.append("qdrant-upsert")
                indexed = SimpleNamespace(publication_id="publication-1", indexed_point_count=1)
                return indexed if finalize is None else finalize(indexed)

        source_id = "source-first-publication"
        path = Path(self.temp_dir.name) / "first-publication.txt"
        path.write_text("first searchable passage", encoding="utf-8")
        self.repository.add_uploaded_knowledge_source(
            source_id, path.name, "TXT", "internal", 0, str(path), path.stat().st_size, "text/plain"
        )

        queue = KnowledgeIngestionQueue(self.repository, index_lifecycle=FreshLifecycle())
        queue.process(source_id, path, "TXT")

        self.assertEqual(events, ["bootstrap-publication", "qdrant-upsert"])
        self.assertEqual(self.repository.get_source_index_status(source_id), "indexed")

    def test_docx_queue_persists_facts_before_qdrant_publication(self) -> None:
        events: list[str] = []
        source_id = "kb-ingestion-facts"
        path = Path(self.temp_dir.name) / "people.docx"
        document = Document()
        document.add_paragraph("姓名：张三，年龄：28岁，性别：女")
        document.save(path)
        self.repository.add_uploaded_knowledge_source(
            source_id,
            "people.docx",
            "文档",
            "公开",
            0,
            str(path),
            path.stat().st_size,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        original_complete = self.repository.complete_knowledge_source_indexing

        def complete_with_fact_event(*args, **kwargs):
            result = original_complete(*args, **kwargs)
            matches = self.repository.find_knowledge_facts(
                WordFactualIntent(
                    entity="张三",
                    entity_normalized=normalize_fact_key("张三"),
                    field="年龄",
                    field_normalized=normalize_fact_key("年龄"),
                )
            )
            if matches:
                events.append("postgres-facts")
            return result

        self.repository.complete_knowledge_source_indexing = complete_with_fact_event

        class RecordingLifecycle:
            def upsert_source(inner_self, source_id, *, finalize=None, on_failure=None):
                events.append("qdrant-upsert")
                return "publication-1", 1

        queue = KnowledgeIngestionQueue(
            self.repository,
            index_lifecycle=RecordingLifecycle(),
        )

        queue.process(source_id, path, "文档")

        self.assertEqual(events, ["postgres-facts", "qdrant-upsert"])

    def test_docx_queue_deduplicates_identical_source_facts_before_persistence(self) -> None:
        source_id = "kb-ingestion-duplicate-facts"
        path = Path(self.temp_dir.name) / "duplicate-people.docx"
        document = Document()
        document.add_paragraph("姓名：张三，年龄：28岁")
        document.add_paragraph("说明" * 400)
        document.add_paragraph("姓名：张三，年龄：28岁")
        document.save(path)
        self.repository.add_uploaded_knowledge_source(
            source_id,
            "duplicate-people.docx",
            "文档",
            "公开",
            0,
            str(path),
            path.stat().st_size,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        self.queue.process(source_id, path, "文档")

        matches = self.repository.find_knowledge_facts(
            WordFactualIntent(
                entity="张三",
                entity_normalized=normalize_fact_key("张三"),
                field="年龄",
                field_normalized=normalize_fact_key("年龄"),
            )
        )
        self.assertEqual(len(matches), 1)
        self.assertEqual(matches[0].fact.locator, {"paragraph": 0})

    def test_qdrant_failure_does_not_delete_postgres_chunks(self) -> None:
        class FailingLifecycle:
            def upsert_source(self, source_id):
                raise RuntimeError("qdrant unavailable")

        source_id = "source-failed-index"
        path = Path(self.temp_dir.name) / "failed-index.txt"
        path.write_text("legacy passage remains available", encoding="utf-8")
        self.repository.add_uploaded_knowledge_source(
            source_id,
            "failed-index.txt",
            "TXT",
            "internal",
            0,
            str(path),
            32,
            "text/plain",
        )
        queue = KnowledgeIngestionQueue(self.repository, index_lifecycle=FailingLifecycle())

        queue.enqueue(source_id, path, "TXT")
        queue.drain()

        self.assertEqual(len(self.repository.list_knowledge_chunks(source_id)), 1)
        self.assertEqual(self.repository.get_source_index_status(source_id), "failed")

    def test_source_deletion_stops_before_postgres_when_qdrant_delete_fails(self) -> None:
        class FailingDeleteLifecycle:
            def delete_source(self, source_id):
                raise RuntimeError("qdrant unavailable")

        source_id = "source-delete"
        self.repository.add_knowledge_source("delete.txt", "TXT", "internal")
        self.repository._state.knowledge_sources[0].id = source_id
        queue = KnowledgeIngestionQueue(
            self.repository,
            index_lifecycle=FailingDeleteLifecycle(),
        )

        with self.assertRaises(KnowledgeIndexUnavailableError) as captured:
            queue.discard_source(source_id)

        self.assertEqual(captured.exception.status_code, 503)
        self.assertTrue(
            any(source.id == source_id for source in self.repository.list_knowledge_sources())
        )

    def test_delete_endpoint_returns_503_and_preserves_source_when_qdrant_fails(self) -> None:
        class FailingDeleteLifecycle:
            def delete_source(self, source_id, *, finalize=None):
                raise RuntimeError("retrieval index reconciliation required")

        source_id = "source-api-delete"
        self.repository.add_knowledge_source("delete.txt", "TXT", "internal")
        self.repository._state.knowledge_sources[0].id = source_id
        queue = KnowledgeIngestionQueue(
            self.repository,
            index_lifecycle=FailingDeleteLifecycle(),
        )
        client = TestClient(
            create_app(
                repository=self.repository,
                upload_dir=Path(self.temp_dir.name),
                ingestion_queue=queue,
            )
        )

        response = client.delete(f"/api/knowledge/sources/{source_id}")

        self.assertEqual(response.status_code, 503)
        self.assertIn("reconciliation", response.json()["detail"])
        self.assertTrue(
            any(source.id == source_id for source in self.repository.list_knowledge_sources())
        )

    def test_delete_endpoint_commits_postgres_inside_source_maintenance_fence(self) -> None:
        events: list[str] = []

        class FencedDeleteLifecycle:
            def delete_source(inner_self, source_id, *, finalize=None):
                events.append("qdrant-delete")
                self.assertIsNotNone(finalize)
                result = finalize()
                events.append("fence-release")
                return result

        source_id = "source-api-fenced-delete"
        self.repository.add_knowledge_source("delete.txt", "TXT", "internal")
        self.repository._state.knowledge_sources[0].id = source_id
        original_delete = self.repository.delete_knowledge_source

        def delete_from_postgres(target_source_id):
            events.append("postgres-delete")
            return original_delete(target_source_id)

        self.repository.delete_knowledge_source = delete_from_postgres
        queue = KnowledgeIngestionQueue(
            self.repository,
            index_lifecycle=FencedDeleteLifecycle(),
        )
        client = TestClient(
            create_app(
                repository=self.repository,
                upload_dir=Path(self.temp_dir.name),
                ingestion_queue=queue,
            )
        )

        response = client.delete(f"/api/knowledge/sources/{source_id}")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(events, ["qdrant-delete", "postgres-delete", "fence-release"])

    def test_reindex_qdrant_failure_preserves_old_postgres_chunks(self) -> None:
        class FailingReindexLifecycle:
            def delete_source(self, source_id, *, finalize=None):
                raise RuntimeError("qdrant unavailable")

        source_id = "source-api-reindex"
        path = Path(self.temp_dir.name) / "reindex.txt"
        path.write_text("replacement content", encoding="utf-8")
        self.repository.add_uploaded_knowledge_source(
            source_id,
            "reindex.txt",
            "TXT",
            "internal",
            0,
            str(path),
            19,
            "text/plain",
        )
        self.repository.complete_knowledge_source_indexing(
            source_id,
            parse_knowledge_file(path, source_id, "TXT"),
        )
        old_chunks = self.repository.list_knowledge_chunks(source_id)
        old_status = next(
            item.status for item in self.repository.list_knowledge_sources() if item.id == source_id
        )
        queue = KnowledgeIngestionQueue(
            self.repository,
            index_lifecycle=FailingReindexLifecycle(),
        )
        client = TestClient(
            create_app(
                repository=self.repository,
                upload_dir=Path(self.temp_dir.name),
                ingestion_queue=queue,
            )
        )

        response = client.post(f"/api/knowledge/sources/{source_id}/reindex")

        self.assertEqual(response.status_code, 503)
        self.assertEqual(self.repository.list_knowledge_chunks(source_id), old_chunks)
        source = next(
            item for item in self.repository.list_knowledge_sources() if item.id == source_id
        )
        self.assertEqual(source.status, old_status)

    def test_reindex_without_active_generation_finalizes_and_requeues(self) -> None:
        events: list[str] = []

        class FreshSystemLifecycle:
            def delete_source(inner_self, source_id, *, finalize=None):
                events.append("safe-qdrant-noop")
                self.assertIsNotNone(finalize)
                retried = finalize()
                events.append("postgres-reindex")
                return retried

            def upsert_source(inner_self, source_id, *, finalize=None, on_failure=None):
                raise AssertionError("reindex should only enqueue during this request")

        source_id = "source-fresh-reindex"
        path = Path(self.temp_dir.name) / "fresh-reindex.txt"
        path.write_text("fresh system replacement", encoding="utf-8")
        self.repository.add_uploaded_knowledge_source(
            source_id,
            "fresh-reindex.txt",
            "TXT",
            "internal",
            0,
            str(path),
            24,
            "text/plain",
        )
        self.repository.complete_knowledge_source_indexing(
            source_id,
            parse_knowledge_file(path, source_id, "TXT"),
        )
        queue = KnowledgeIngestionQueue(
            self.repository,
            index_lifecycle=FreshSystemLifecycle(),
        )
        client = TestClient(
            create_app(
                repository=self.repository,
                upload_dir=Path(self.temp_dir.name),
                ingestion_queue=queue,
            )
        )

        response = client.post(f"/api/knowledge/sources/{source_id}/reindex")

        self.assertEqual(response.status_code, 200, response.text)
        self.assertEqual(events, ["safe-qdrant-noop", "postgres-reindex"])
        source = next(item for item in response.json() if item["id"] == source_id)
        self.assertEqual(source["records"], 0)


if __name__ == "__main__":
    unittest.main()
