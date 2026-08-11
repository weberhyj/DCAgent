from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.docx_parser import parse_docx_knowledge_file, read_docx_blocks
from app.text_parser import parse_knowledge_file
from app.word_facts import KnowledgeFactModel, normalize_fact_key


ParagraphFixture = str | tuple[str, str]


def write_docx(path: Path, paragraphs: list[ParagraphFixture]) -> Path:
    document = Document()
    for fixture in paragraphs:
        if isinstance(fixture, tuple):
            text, style = fixture
            document.add_paragraph(text, style=style)
        else:
            document.add_paragraph(fixture)
    document.save(path)
    return path


def write_docx_table(path: Path, headers: list[str], rows: list[list[str]]) -> Path:
    document = Document()
    table = document.add_table(rows=1, cols=len(headers))
    for column, value in enumerate(headers):
        table.rows[0].cells[column].text = value
    for values in rows:
        cells = table.add_row().cells
        for column, value in enumerate(values):
            cells[column].text = value
    document.save(path)
    return path


class DocxParserTest(unittest.TestCase):
    def setUp(self) -> None:
        self.temp_dir = Path(tempfile.mkdtemp())

    def tearDown(self) -> None:
        for path in self.temp_dir.iterdir():
            path.unlink()
        self.temp_dir.rmdir()

    def test_inline_record_extracts_only_known_fields(self) -> None:
        path = write_docx(
            self.temp_dir / "people.docx",
            paragraphs=["姓名：张三，年龄：28岁，性别：女，职务：工程师，爱好：登山"],
        )

        result = parse_docx_knowledge_file(path, source_id="kb-people")

        self.assertEqual(
            {(item.entity, item.field, item.value) for item in result.facts},
            {
                ("张三", "年龄", "28岁"),
                ("张三", "性别", "女"),
                ("张三", "职务", "工程师"),
            },
        )
        self.assertEqual({item.confidence for item in result.facts}, {0.97})

    def test_extracted_fields_pass_the_shared_canonical_contract(self) -> None:
        path = write_docx(
            self.temp_dir / "shared-contract.docx",
            paragraphs=["姓名：张三，岁数：28岁，男女：女，职位：工程师"],
        )

        result = parse_docx_knowledge_file(path, source_id="kb-shared-contract")

        self.assertEqual([fact.field for fact in result.facts], ["年龄", "性别", "职务"])
        for index, fact in enumerate(result.facts):
            recreated = KnowledgeFactModel.create(
                id=f"fact-recreated-{index}",
                source_id=fact.source_id,
                chunk_id=fact.chunk_id,
                entity=fact.entity,
                field=fact.field,
                value=fact.value,
                confidence=fact.confidence,
                locator=fact.locator,
            )
            self.assertEqual(recreated.field, fact.field)
            self.assertEqual(fact.field_normalized, normalize_fact_key(fact.field))

    def test_table_row_extracts_fields_and_locator(self) -> None:
        path = write_docx_table(
            self.temp_dir / "staff.docx",
            headers=["姓名", "年龄", "性别", "职务"],
            rows=[["李四", "31岁", "男", "产品经理"]],
        )

        result = parse_docx_knowledge_file(path, source_id="kb-staff")

        age = next(item for item in result.facts if item.field == "年龄")
        self.assertEqual(age.entity, "李四")
        self.assertEqual(age.locator, {"table": 0, "row": 1, "column": 1})
        self.assertEqual(age.confidence, 0.99)

    def test_heading_entity_applies_to_following_key_value_lines(self) -> None:
        path = write_docx(
            self.temp_dir / "heading.docx",
            paragraphs=[("张三", "Heading 1"), "年龄：28岁", "性别：女"],
        )

        result = parse_docx_knowledge_file(path, source_id="kb-heading")

        self.assertEqual(
            {(item.field, item.value) for item in result.facts},
            {("年龄", "28岁"), ("性别", "女")},
        )
        self.assertEqual({item.entity for item in result.facts}, {"张三"})
        self.assertEqual({item.confidence for item in result.facts}, {0.95})

    def test_next_heading_stops_the_previous_entity(self) -> None:
        path = write_docx(
            self.temp_dir / "two-headings.docx",
            paragraphs=[
                ("张三", "Heading 1"),
                "年龄：28岁",
                ("李四", "Heading 1"),
                "年龄：31岁",
            ],
        )

        result = parse_docx_knowledge_file(path, source_id="kb-headings")

        self.assertEqual(
            {(item.entity, item.field, item.value) for item in result.facts},
            {("张三", "年龄", "28岁"), ("李四", "年龄", "31岁")},
        )

    def test_narrative_sentence_remains_chunk_only(self) -> None:
        path = write_docx(
            self.temp_dir / "narrative.docx",
            paragraphs=["张三是一名工程师，今年二十八岁，性格开朗。"],
        )

        result = parse_docx_knowledge_file(path, source_id="kb-narrative")

        self.assertEqual(result.facts, ())
        self.assertIn("二十八岁", "\n".join(item.text for item in result.chunks))

    def test_table_requires_exactly_one_entity_header(self) -> None:
        path = write_docx_table(
            self.temp_dir / "ambiguous-table.docx",
            headers=["姓名", "员工", "年龄"],
            rows=[["张三", "E-001", "28岁"]],
        )

        result = parse_docx_knowledge_file(path, source_id="kb-ambiguous")

        self.assertEqual(result.facts, ())

    def test_blocks_keep_paragraph_and_table_document_order(self) -> None:
        path = self.temp_dir / "ordered.docx"
        document = Document()
        document.add_paragraph("表格之前")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "姓名"
        table.rows[0].cells[1].text = "年龄"
        row = table.add_row().cells
        row[0].text = "张三"
        row[1].text = "28岁"
        document.add_paragraph("表格之后")
        document.save(path)

        blocks = read_docx_blocks(path)

        self.assertEqual(
            [block.text for block in blocks],
            ["表格之前", "姓名 | 年龄", "张三 | 28岁", "表格之后"],
        )
        self.assertEqual(blocks[0].locator, {"paragraph": 0})
        self.assertEqual(blocks[2].locator, {"table": 0, "row": 1})
        self.assertEqual(blocks[3].locator, {"paragraph": 1})

    def test_facts_follow_paragraph_and_table_document_order(self) -> None:
        path = self.temp_dir / "fact-order.docx"
        document = Document()
        document.add_paragraph("姓名：张三，年龄：28岁")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "姓名"
        table.rows[0].cells[1].text = "年龄"
        row = table.add_row().cells
        row[0].text = "李四"
        row[1].text = "31岁"
        document.add_paragraph("姓名：王五，年龄：42岁")
        document.save(path)

        result = parse_docx_knowledge_file(path, source_id="kb-fact-order")

        self.assertEqual(
            [(fact.entity, fact.value) for fact in result.facts],
            [("张三", "28岁"), ("李四", "31岁"), ("王五", "42岁")],
        )

    def test_long_single_block_uses_existing_overlap_and_keeps_locator(self) -> None:
        text = "".join(str(index % 10) for index in range(720))
        path = write_docx(self.temp_dir / "long.docx", paragraphs=[text])

        result = parse_docx_knowledge_file(path, source_id="kb-long")

        self.assertEqual(len(result.chunks), 2)
        self.assertEqual(result.chunks[0].text[-120:], result.chunks[1].text[:120])
        self.assertEqual(
            [chunk.metadata["locators"] for chunk in result.chunks],
            [[{"paragraph": 0}], [{"paragraph": 0}]],
        )

    def test_oversized_inline_fact_references_chunk_containing_late_value(self) -> None:
        path = write_docx(
            self.temp_dir / "long-inline-fact.docx",
            paragraphs=[f"姓名：张三，备注：{'x' * 650}，年龄：73岁"],
        )

        result = parse_docx_knowledge_file(path, source_id="kb-long-inline")

        fact = next(item for item in result.facts if item.field == "年龄")
        evidence_chunk = next(chunk for chunk in result.chunks if chunk.id == fact.chunk_id)
        self.assertIn(fact.value, evidence_chunk.text)

    def test_oversized_table_fact_references_chunk_containing_cell_value(self) -> None:
        path = write_docx_table(
            self.temp_dir / "long-table-fact.docx",
            headers=["姓名", "备注", "年龄"],
            rows=[["李四", "x" * 650, "74岁"]],
        )

        result = parse_docx_knowledge_file(path, source_id="kb-long-table")

        fact = next(item for item in result.facts if item.field == "年龄")
        evidence_chunk = next(chunk for chunk in result.chunks if chunk.id == fact.chunk_id)
        self.assertIn(fact.value, evidence_chunk.text)

    def test_table_row_stays_whole_and_preserves_cell_locators(self) -> None:
        path = self.temp_dir / "row-locator.docx"
        document = Document()
        document.add_paragraph("前置内容" * 115)
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "姓名"
        table.rows[0].cells[1].text = "年龄"
        row = table.add_row().cells
        row[0].text = "王五"
        row[1].text = "42岁"
        document.save(path)

        result = parse_docx_knowledge_file(path, source_id="kb-row")

        row_chunk = next(chunk for chunk in result.chunks if "王五 | 42岁" in chunk.text)
        self.assertIn({"table": 0, "row": 1, "column": 0}, row_chunk.metadata["locators"])
        self.assertIn({"table": 0, "row": 1, "column": 1}, row_chunk.metadata["locators"])
        self.assertIn("王五 | 42岁", row_chunk.text)

    def test_fact_ids_are_stable_across_repeated_parses(self) -> None:
        path = write_docx(
            self.temp_dir / "stable.docx",
            paragraphs=["姓名：张三，年龄：28岁"],
        )

        first = parse_docx_knowledge_file(path, source_id="kb-stable")
        second = parse_docx_knowledge_file(path, source_id="kb-stable")

        self.assertEqual([fact.id for fact in first.facts], [fact.id for fact in second.facts])
        self.assertEqual(
            [chunk.id for chunk in first.chunks],
            [chunk.id for chunk in second.chunks],
        )

    def test_legacy_parse_function_still_returns_list_of_chunks(self) -> None:
        path = write_docx(
            self.temp_dir / "compat.docx",
            paragraphs=["姓名：张三，年龄：28岁"],
        )

        chunks = parse_knowledge_file(path, "kb-compat", "文档")

        self.assertIsInstance(chunks, list)
        self.assertGreater(len(chunks), 0)


if __name__ == "__main__":
    unittest.main()
