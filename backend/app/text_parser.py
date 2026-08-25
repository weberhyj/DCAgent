from __future__ import annotations

import csv
import re
from io import StringIO
from pathlib import Path
from uuid import uuid4

from .docx_parser import KnowledgeParseResult, parse_docx_knowledge_file, split_text_spans
from .models import KnowledgeChunkModel

CHUNK_SIZE = 600
CHUNK_OVERLAP = 120


def parse_knowledge_file(path: Path, source_id: str, source_type: str) -> list[KnowledgeChunkModel]:
    return list(parse_knowledge_file_result(path, source_id, source_type).chunks)


def parse_knowledge_file_result(
    path: Path,
    source_id: str,
    source_type: str,
) -> KnowledgeParseResult:
    if path.suffix.lower() == ".docx":
        return parse_docx_knowledge_file(path, source_id)
    suffix = path.suffix.lower()
    extracted_text = extract_text(path, source_type)
    chunker = chunk_tabular_text if suffix in {".csv", ".xlsx"} else chunk_text
    return KnowledgeParseResult(
        chunks=tuple(chunker(source_id, extracted_text)),
        facts=(),
    )


def extract_text(path: Path, source_type: str) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return read_text_with_fallback(path)
    if suffix == ".csv":
        return extract_csv_text(path)
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".xlsx":
        return extract_xlsx_text(path)
    return read_binary_as_text(path)


def read_text_with_fallback(path: Path) -> str:
    content = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def read_binary_as_text(path: Path) -> str:
    return path.read_bytes().decode("utf-8", errors="ignore")


def extract_csv_text(path: Path) -> str:
    raw = read_text_with_fallback(path)
    rows = csv.reader(StringIO(raw))
    rendered_rows = [_render_tabular_row(row) for row in rows]
    rendered_rows = [row for row in rendered_rows if row]
    return "\n".join(["[CSV]", *rendered_rows]) if rendered_rows else ""


def extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception:
        return read_binary_as_text(path)

    try:
        reader = PdfReader(str(path))
        pages = [(page.extract_text() or "") for page in reader.pages]
        return "\n".join(pages)
    except Exception:
        return read_binary_as_text(path)


def extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except Exception:
        return read_binary_as_text(path)

    try:
        document = Document(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_rows = []
        for table in document.tables:
            for row in table.rows:
                rendered_row = _render_tabular_row(cell.text for cell in row.cells)
                if rendered_row:
                    table_rows.append(rendered_row)
        return "\n".join([*paragraphs, *table_rows])
    except Exception:
        return read_binary_as_text(path)


def extract_xlsx_text(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except Exception:
        return read_binary_as_text(path)

    try:
        workbook = load_workbook(str(path), read_only=True, data_only=True)
        rows: list[str] = []
        for sheet in workbook.worksheets:
            rows.append(f"[{sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                rendered_row = _render_tabular_row(row)
                if rendered_row:
                    rows.append(rendered_row)
        workbook.close()
        return "\n".join(rows)
    except Exception:
        return read_binary_as_text(path)


def normalize_text(text: str) -> str:
    normalized = text.replace("\x00", "")
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def chunk_text(source_id: str, text: str) -> list[KnowledgeChunkModel]:
    normalized = normalize_text(text)
    if not normalized:
        normalized = "空白文件"

    chunks: list[KnowledgeChunkModel] = []
    for start, end in split_text_spans(
        normalized,
        max_length=CHUNK_SIZE,
        overlap=CHUNK_OVERLAP,
    ):
        chunk = normalized[start:end].strip()
        if chunk:
            chunks.append(
                KnowledgeChunkModel(
                    id=f"chunk-{uuid4().hex[:10]}",
                    source_id=source_id,
                    chunk_index=len(chunks),
                    text=chunk,
                    token_count=estimate_token_count(chunk),
                )
            )
    return chunks


def chunk_tabular_text(source_id: str, text: str) -> list[KnowledgeChunkModel]:
    """Chunk CSV/XLSX text by complete rows and repeat table context.

    The first non-empty row in each sheet/CSV section is treated as the
    header.  Every emitted chunk repeats the sheet marker and header so a
    retrieved data row remains understandable without relying on a previous
    chunk.  A single oversized row is split only as a last resort, using the
    same semantic splitter as narrative text.
    """

    normalized = normalize_text(text)
    if not normalized:
        normalized = "空白文件"

    sections = _tabular_sections(normalized)
    chunks: list[KnowledgeChunkModel] = []

    def append_chunk(chunk_text_value: str) -> None:
        clean_text = chunk_text_value.strip()
        if not clean_text:
            return
        chunks.append(
            KnowledgeChunkModel(
                id=f"chunk-{uuid4().hex[:10]}",
                source_id=source_id,
                chunk_index=len(chunks),
                text=clean_text,
                token_count=estimate_token_count(clean_text),
            )
        )

    for label, rows in sections:
        if not rows:
            if label:
                append_chunk(label)
            continue

        # Most workbooks/CSVs have a header, but a user-provided export can
        # contain data immediately on the first row.  Treat an obviously
        # data-shaped first row as a headerless table instead of silently
        # dropping that row from ``data_rows``.  This is deliberately
        # conservative: a normal textual header remains the header, while a
        # row containing dates/numbers is retained as data with a synthetic
        # context label.
        has_header = _looks_like_tabular_header(rows[0])
        header = rows[0] if has_header else "数据行（无表头）"
        context = _tabular_context(label, header)
        data_rows = rows[1:] if has_header else rows
        if not data_rows:
            append_chunk(context)
            continue

        pending_rows: list[str] = []

        def flush_rows(rows_to_flush: list[str], context_text: str) -> None:
            if rows_to_flush:
                append_chunk("\n".join([context_text, *rows_to_flush]))
                rows_to_flush.clear()

        for row in data_rows:
            available = CHUNK_SIZE - len(context) - 1
            if len(row) <= available:
                candidate_length = len(context) + sum(len(item) for item in pending_rows)
                candidate_length += len(pending_rows) + len(row) + 1
                if pending_rows and candidate_length > CHUNK_SIZE:
                    flush_rows(pending_rows, context)
                pending_rows.append(row)
                continue

            flush_rows(pending_rows, context)
            # A nearly full row cannot carry a header prefix without being
            # split.  Preserve the complete row and sacrifice the optional
            # context; the row itself remains searchable and column order is
            # still intact.
            if len(row) <= CHUNK_SIZE and len(row) >= CHUNK_SIZE - 1:
                append_chunk(row)
                continue
            if len(row) < CHUNK_SIZE:
                row_context_limit = CHUNK_SIZE - len(row) - 1
                row_context = _tabular_context(
                    label,
                    header,
                    context_limit=max(0, row_context_limit),
                )
                if row_context and len(row_context) + len(row) + 1 <= CHUNK_SIZE:
                    append_chunk("\n".join([row_context, row]))
                    continue
            if available <= 0:
                # A wide sheet/header can consume the whole chunk budget. Do
                # not prepend that context to a full-size row slice, or the
                # resulting chunk would exceed CHUNK_SIZE and lose the very
                # row that needs to remain searchable. The bounded context
                # chunk above still preserves the header as a standalone
                # hint; row slices must use the complete budget here.
                append_chunk(context)
                row_available = CHUNK_SIZE
                row_context = ""
            else:
                row_available = available
                row_context = context
            for start, end in split_text_spans(
                row,
                max_length=max(1, row_available),
                overlap=min(CHUNK_OVERLAP, max(1, row_available // 2)),
            ):
                append_chunk(
                    "\n".join(
                        part
                        for part in (row_context, row[start:end])
                        if part
                    )
                )

        flush_rows(pending_rows, context)

    if not chunks:
        append_chunk("空白文件")
    return chunks


def _tabular_sections(text: str) -> list[tuple[str, list[str]]]:
    sections: list[tuple[str, list[str]]] = []
    label = ""
    rows: list[str] = []
    for line in text.splitlines():
        clean_line = line.strip()
        if not clean_line:
            continue
        if clean_line.startswith("[") and clean_line.endswith("]"):
            if label or rows:
                sections.append((label, rows))
            label = clean_line
            rows = []
            continue
        rows.append(clean_line)
    if label or rows:
        sections.append((label, rows))
    return sections


def _render_tabular_row(values: object) -> str:
    """Render a row without shifting values when an interior cell is blank."""

    if isinstance(values, (str, bytes)):
        cells = [values]
    else:
        try:
            cells = list(values)  # type: ignore[arg-type]
        except TypeError:
            cells = [values]

    rendered = []
    for value in cells:
        if value is None:
            rendered.append("")
            continue
        rendered.append(
            str(value).strip().replace("\r", " ").replace("\n", " ")
        )
    while rendered and not rendered[-1]:
        rendered.pop()
    return " | ".join(rendered)


_TABULAR_NUMERIC_CELL_RE = re.compile(
    r"^[+-]?(?:\d+(?:\.\d+)?|\.\d+)%?$"
)
_TABULAR_DATE_CELL_RE = re.compile(
    r"^(?:\d{4}[-/]\d{1,2}[-/]\d{1,2}(?:[T ]\d{1,2}:\d{2}(?::\d{2}(?:\.\d+)?)?)?"
    r"|\d{1,2}:\d{2}(?::\d{2})?)$"
)


def _looks_like_tabular_header(row: str) -> bool:
    """Return whether a rendered row is likely a field-header row.

    This is only a loss-prevention heuristic for the unstructured chunker;
    structured Excel ingestion still performs its own schema inference and
    confirmation.  A row is considered data-shaped when at least half of its
    cells are numeric/date-like, or when it contains a long free-text value.
    """

    cells = [cell.strip() for cell in row.split(" | ") if cell.strip()]
    if not cells:
        return False
    data_like = 0
    for cell in cells:
        compact = re.sub(r"\s+", "", cell)
        if _TABULAR_NUMERIC_CELL_RE.fullmatch(compact) or _TABULAR_DATE_CELL_RE.fullmatch(
            compact
        ):
            data_like += 1
            continue
        # Long narrative values are much more likely to be a record than a
        # field label.  Keep the threshold high to avoid classifying verbose
        # headers such as ``客户详细地址`` as data accidentally.
        if len(cell) > 48:
            data_like += 1
    return data_like < max(1, (len(cells) + 1) // 2)


def _tabular_context(
    label: str,
    header: str,
    *,
    context_limit: int = 300,
) -> str:
    """Build a bounded, self-describing context prefix for each data row."""

    lines = [line for line in (label, header) if line]
    context = "\n".join(lines)
    # Reserve room for at least a useful portion of a row.  Wide spreadsheets
    # can have thousands of characters in a generated header; retaining the
    # beginning of the sheet/header is preferable to producing chunks larger
    # than the ingestion contract or dropping every data row.
    context_limit = min(context_limit, max(0, CHUNK_SIZE - 1))
    if context_limit <= 0:
        return ""
    if len(context) <= context_limit:
        return context

    if label and header:
        if context_limit <= 2:
            return _bounded_text(label or header, context_limit)
        label_budget = min(len(label), 80, max(1, (context_limit - 1) // 2))
        header_budget = max(1, context_limit - label_budget - 1)
        return f"{_bounded_text(label, label_budget)}\n{_bounded_text(header, header_budget)}"
    return _bounded_text(context, context_limit)


def _bounded_text(value: str, limit: int) -> str:
    if limit <= 0:
        return ""
    if len(value) <= limit:
        return value
    if limit <= 2:
        return value[:limit]
    head_length = (limit - 1) // 2
    tail_length = limit - head_length - 1
    return f"{value[:head_length]}…{value[-tail_length:]}"


def estimate_token_count(text: str) -> int:
    ascii_words = re.findall(r"[A-Za-z0-9_]+", text)
    non_ascii_chars = [char for char in text if ord(char) > 127 and not char.isspace()]
    return max(1, len(ascii_words) + len(non_ascii_chars))
